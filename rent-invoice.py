import docx
from tkinter import *
import os

pwd = os.getcwd()
import sys


def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    base_path = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base_path, relative_path)


root = Tk()

root.title('Rent Invoice Automation')
root.state('zoomed')
try:
    root.wm_iconbitmap("C:\\Users\\darsh\\PycharmProjects\\test111\\test11\\New folder\\homelogo.ico")
    try:
        icon_path = resource_path('C:\\Users\\darsh\\PycharmProjects\\test111\\test11\\New folder\\homelogo.ico')
        root.iconbitmap(icon_path)
    except:
        pass
except:
    pass

image_num = 0


try:
    background_image = PhotoImage(file="C:\\Users\\darsh\\PycharmProjects\\test111\\test11\\New folder\\landscape.png")
    background_label = Label(root,image=background_image)
    background_label.place(x=0,y=0)
except:
    pass


def background():
    try:
        global background_label
        background_label.destroy()
        backgroundbutton.destroy()

    except:
        pass


def myfunction(event):
    canvas.configure(scrollregion=canvas.bbox("all"), width=700, height=400)


myframe = Frame(root, bd=5, width=700, height=600, relief=SUNKEN, highlightthickness=5, highlightcolor='black',
                highlightbackground='white')
myframe.place(x=300, y=0)

canvas = Canvas(myframe)
frame = Frame(canvas)

myscrollbar = Scrollbar(myframe, orient="vertical", command=canvas.yview)
canvas.configure(yscrollcommand=myscrollbar.set)
myscrollbar.pack(side="right", fill="y")
canvas.pack(side="right")
canvas.create_window((0, 0), window=frame, anchor='nw')
frame.bind("<Configure>", myfunction)

from tkinter import messagebox

import sqlite3

conn = sqlite3.connect("rent-invoice.db")
cur = conn.cursor()
# cur.execute('SELECT COUNT(*) FROM MY_TABLE;')
# exist = cur.fetchone()

try:
    column_sql = cur.execute("SELECT * from 'invoices' limit 1")
    col_name = [i[0] for i in column_sql.description]
    # print(col_name)
except:
    cur.execute('''
               CREATE TABLE IF NOT EXISTS invoices
               (Date varchar(250), 
               `Invoice Month-Year` varchar(250),
               Address varchar(250) ,
               Rent decimal (50,4),
               `Pet fees` decimal (50,4),
               `Water and Trash Fees` decimal (50,4),
               `Sewerage fees` decimal (50,4),
               `Electricity Charges` decimal (50,4),
               `Other Fees` decimal (50,4),
               Total decimal (50,4)
               )''')
    messagebox.showinfo("Welcome",
                        'The extra expenses you add will be saved for future inputs. To reset, reinstall the app')

expenses = ['Date', 'Invoice Month-Year', 'Address', 'Rent', 'Pet fees', 'Water and Trash Fees', 'Sewerage fees',
            'Electricity Charges', 'Other Fees']

expenses_update = expenses
try:
    
    if len(col_name) != (len(expenses) + 1):
        default_len = len(expenses)
        new_len = len(col_name) - 1
        extras = new_len - default_len
        # print(extras)
        for extra in range(extras):
            num = default_len + extra + 1
            # print(num)
            if col_name[num] != 'Total':
                expenses_update.append(col_name[num])
                
except:
    pass

i = 2
labels = {}
newlabel_entry = {}
k = 0
for expense in expenses_update:
    labels[k] = Label(frame, text=expense)
    labels[k].grid(row=i, column=1)
    newlabel_entry[k] = Entry(frame, justify=CENTER, width=40, bd=4)
    newlabel_entry[k].grid(row=i, column=2)
    k += 1
    i += 1

row_number_for_saving_file = i
input_savelabel = Label(frame, text='New file name without extension (Existing file might be overwritten)')
input_savelabel.grid(row=i, column=1)
input_save = Entry(frame, justify=CENTER, width=40, bd=4)
input_save.grid(row=i, column=2)
i += 1

abcde = None
name_newlabel = {}
amount_newlabel = {}
name_new = {}
amount_new = {}
a = 0


def funct2():
    global i
    global a
    i += 1
    a += 1
    name_newlabel[a] = Label(frame, text='Expense name')
    name_newlabel[a].grid(row=i, column=1)
    name_new[a] = Entry(frame, justify=CENTER, width=40, bd=4)
    name_new[a].grid(row=i, column=2)
    i += 1
    amount_newlabel[a] = Label(frame, text='Expense amount')
    amount_newlabel[a].grid(row=i, column=1)
    amount_new[a] = Entry(frame, justify=CENTER, width=40, bd=4)
    amount_new[a].grid(row=i, column=2)


addbutton = Button(frame, text='Add', activebackground='#00ff00', command=funct2, bd=4)
addbutton.grid(row=i, column=2)


def delfunction():
    global a
    try:
        name_newlabel[a].destroy()
        name_new[a].destroy()
        amount_newlabel[a].destroy()
        amount_new[a].destroy()
        a -= 1
    except:
        pass


delbutton = Button(frame, text='Delete', activebackground='#00ff00', command=delfunction, justify=RIGHT, bd=4)
delbutton.grid(row=i, column=3)


def func1():
    global a
    global label123
    global file_exist_and_open

    abcdef = []

    try:
        for j in range(1, a + 1):
            abc = name_new[j].get()
            abcd = amount_new[j].get()
            abcde = [abc, abcd]
            abcdef.append(abcde)
    except:
        pass
    global k
    all_expenses = []
    for m in range(3, k):
        expense_amount_entered = newlabel_entry[m].get()
        expenses_label_list = [labels[m].cget("text"), expense_amount_entered]
        all_expenses.append(expenses_label_list)
    date_get = newlabel_entry[0].get()
    global inv_get
    inv_get = newlabel_entry[1].get()
    add_get = newlabel_entry[2].get()
    input_save1 = input_save.get()

    doc = docx.Document()
    doc.add_heading('Date: ' + date_get, 5)
    doc.add_heading('Invoice Month:' + inv_get, 5)
    doc.add_heading('Property Address: ' + add_get, 5)
    doc.add_paragraph()
    doc.add_paragraph()
    records = all_expenses
        if abcdef != []:
        for j in abcdef:
            records.append(j)
         
    try:
        expense_values = [float(i[1]) for i in records]
        total = sum(expense_values)
        table1 = doc.add_table(rows=0, cols=2)
        table1.style = 'Table Grid'
        for reason, amount in records:
            row_Cells = table1.add_row().cells
            row_Cells[0].text = reason
            row_Cells[1].text = str(amount)
        row_Cells = table1.add_row().cells
        row_Cells[0].text = 'Total'
        row_Cells[1].text = str(total)
        doc_name = input_save1 + '.docx'
        folders = os.listdir()
        try:
            if "Invoices" not in folders:
                os.mkdir('Invoices')
            doc.save('Invoices\\' + doc_name)
        except OSError:
            doc.save(doc_name)
        label_date = Label(frame, text=date_get)
        label_date.grid(row=2, column=2)
        label_inv = Label(frame, text=inv_get)
        label_inv.grid(row=3, column=2)
        label_add = Label(frame, text=add_get)
        label_add.grid(row=4, column=2)
        row_num = 5
        for entry_to_label in records:
            values = entry_to_label[1]
            new_entries = Label(frame, text=values)
            new_entries.grid(row=row_num, column=2)
            row_num += 1
        for entry in range(0, len(newlabel_entry)):
            newlabel_entry[entry].destroy()
        newlabel_entry[0].destroy()
        newlabel_entry[1].destroy()
        newlabel_entry[2].destroy()
        input_save.destroy()

        global filename
        filename = input_save1
        labelinput_save = Label(frame, text=input_save1)
        labelinput_save.grid(row=row_number_for_saving_file, column=2)
        addbutton.destroy()
        delbutton.destroy()
        global i
        labelname_new = {}
        labelamount_new = {}

        for j in range(0, a):
            i += 1
            labelname_new[j] = Label(frame, text=abcdef[j][0])
            labelname_new[j].grid(row=i, column=1)
            labelamount_new[j] = Label(frame, text=abcdef[j][1])
            labelamount_new[j].grid(row=i, column=2)
            name_new[j + 1].destroy()
            amount_new[j + 1].destroy()
            name_newlabel[j + 1].destroy()
            amount_newlabel[j + 1].destroy()
        messagebox.showinfo('Success', 'Word file created in the Invoices directory')
        # successlabel = Label(frame, text='Word file created in Invoices in same directory', fg='green')
        # successlabel.grid(row=1, column=2)
        global uploadondrive
        uploadondrive = Button(frame, text='Upload on drive', activebackground='#00ff00', command=googledrive_entries)
        uploadondrive.grid(row=1, column=1)
        button1.destroy()
        try:
            # integrating with DB
            cur.execute(
                "INSERT INTO invoices('Date', 'Invoice Month-Year', 'Address') VALUES(?,?,?)",
                (date_get, inv_get, add_get))
            for record in records:
                if record[0] not in expenses_update:
                    query = "ALTER TABLE invoices ADD COLUMN `" + record[0] + "`decimal (50,4)"
                    cur.execute(query)

            for record in records:
                query = "UPDATE invoices SET `" + str(
                    record[0]) + "` = (?) WHERE (Date=(?) AND Address=(?)) AND `Invoice Month-Year`=(?)"
                # print(query)
                cur.execute(query, (record[1], date_get, add_get, inv_get))
            cur.execute("UPDATE invoices SET Total = (?) WHERE (Date=(?) AND Address=(?)) AND `Invoice Month-Year`=(?)",
                        (total, date_get, add_get, inv_get))
            conn.commit()
        except:
            messagebox.showerror('Error',
                                 'Invoice created but failed to upload on Database. The file can still be uploaded on drive.')

        return doc_name
    except ValueError:
        messagebox.showerror('Error', 'Enter amount in integers or decimals only')
        
    except OSError:
        messagebox.showerror('Error',
                             'Word file name you mentioned is already open. Close and submit to update the same file.')


button1 = Button(frame, text='Submit', bd=4, activebackground='#00ff00', command=func1)
button1.grid(row=1, column=1)


def googledrive_entries():
    global i
    folder_namelabel = Label(frame, text='Enter the folder name you want to create/present on drive: ')
    folder_namelabel.grid(row=i + 2, column=1)
    global folder_name
    folder_name = Entry(frame, justify=CENTER, width=40, bd=4)
    folder_name.grid(row=i + 2, column=2)
    file_namelabel = Label(frame, text='Enter the file name for the word document you want to give on drive: ')
    file_namelabel.grid(row=i + 3, column=1)
    global title
    title = Entry(frame, justify=CENTER, width=40, bd=4)
    title.grid(row=i + 3, column=2)

    fullpathlabel = Label(frame, text='Do not change path if file not moved ')
    fullpathlabel.grid(row=i + 4, column=1)
    global fullpath
    fullpath = Entry(frame, justify=CENTER, width=40, bd=4)
    fullpath.grid(row=i + 4, column=2)
    global filename, pwd
    fullpath.delete(0, END)
    pwd2slash = pwd.split('\\')
    path = ''
    for j in pwd2slash[:]:
        path = path + j + '\\'
    finalpath = path + 'Invoices\\' + filename + '.docx'
    fullpath.insert(0, finalpath)
    submitbutton.grid(row=i + 5, column=1)
    i += 5
    global uploadondrive
    uploadondrive.destroy()




def check_folder_exists_and_create_fileupload(folder_name1, title1, fullpath1):
    from pydrive.auth import GoogleAuth
    from pydrive.drive import GoogleDrive
    GoogleAuth.DEFAULT_SETTINGS[
        'client_config_file'] = "C:\\Users\\darsh\\PycharmProjects\\test111\\test11\\New folder\\client_secrets.json"

    
    gauth = GoogleAuth()
    
    drive = GoogleDrive(gauth)
    global i

    def check_folder_exists(folder_name):
        list_of_file = drive.ListFile({'q': "'root' in parents and trashed=false"}).GetList()
        global folder_id
        list1 = []
        for drive_folder in list_of_file:
            list1.append(drive_folder['title'])
        if folder_name in list1:
            for drive_folder in list_of_file:
                if drive_folder['title'] == folder_name:
                    folder_id = drive_folder['id']
                    return folder_id

        else:
            folder = drive.CreateFile({'title': folder_name, "mimeType": "application/vnd.google-apps.folder"})
            folder.Upload()
            folder_id = folder['id']
            return folder_id

    fid1 = check_folder_exists(folder_name1)

    def upload_file_inside_folder(title, fid1, fullpath):
        file = drive.CreateFile({'title': title, 'parents': [{'kind': 'drive#fileLink', 'id': fid1}]})
        file.SetContentFile(fullpath)
        file.Upload()
        return True

    global successdrive
    try:
        upload_file_inside_folder(title1, fid1, fullpath1)
        messagebox.showinfo('Success', 'File successfully uploaded on drive.')
        sendmail.grid(column=1)

    except Exception:
        messagebox.showerror('Error', 'Path Error in uploading on drive')


submitbutton = Button(frame, text='Submit on drive', activebackground='#00ff00',
                      command=lambda: check_folder_exists_and_create_fileupload(folder_name.get(), title.get(),
                                                                                fullpath.get()))


def sendmailfunction():
    sendmail.destroy()
    global fromadd, password, to
    global i
    tolabel = Label(frame, text='Send to')
    tolabel.grid(row=i + 2, column=1)
    to = Entry(frame, justify=CENTER, width=40, bd=4)
    to.grid(row=i + 2, column=2)
    submitmail.grid(row=i + 5, column=1)
    i += 2




def bindmailelementsusingAPI(to1):
    from Google import Create_Service
    import base64
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    global i
    global folder_id
    CLIENT_SECRET_FILE = "C:\\Users\\darsh\\PycharmProjects\\test111\\test11\\New folder\\client_secrets.json"
    API_NAME = 'gmail'
    API_VERSION = 'v1'
    SCOPES = ['https://mail.google.com/']
    global authorizeerror, successsent, sendingerror

    try:
        service = Create_Service(CLIENT_SECRET_FILE, API_NAME, API_VERSION, SCOPES)

    except:

        messagebox.showerror('Error', 'Authorization file moved')
        
    try:
        totenant = to1
        emailMsg = ''' Hi, 

        Your monthly rent invoice has been uploaded on the drive. 
        https://drive.google.com/drive/folders/''' + str(folder_id) + '''

        Thanks. '''
        global inv_get
        mimeMessage = MIMEMultipart()
        mimeMessage['to'] = totenant
        mimeMessage['subject'] = 'Rent Invoice ' + inv_get + ' Uploaded on Drive'
        mimeMessage.attach(MIMEText(emailMsg, 'plain'))
        raw_string = base64.urlsafe_b64encode(mimeMessage.as_bytes()).decode()

        message = service.users().messages().send(userId='me', body={'raw': raw_string}).execute()
        # print(message)
        messagebox.showinfo('Success', 'Mail successfully sent')
  

    except:
       
        messagebox.showinfo('Error', 'Error in sending email')
    


sendmail = Button(frame, text='Send Notification Email', activebackground='#00ff00', command=sendmailfunction)


submitmail = Button(frame, text='Send email', activebackground='#00ff00',
                    command=lambda: bindmailelementsusingAPI(to.get()))

root.mainloop()


